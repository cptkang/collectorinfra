"""Word 양식 구조 분석 모듈.

python-docx를 사용하여 Word 파일의 플레이스홀더, 표 구조, 스타일 정보를 추출한다.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any

from docx import Document

logger = logging.getLogger(__name__)

# {{placeholder}} 패턴 정규식
_PLACEHOLDER_PATTERN = re.compile(r"\{\{(.+?)\}\}")


def parse_word_template(file_data: bytes) -> dict[str, Any]:
    """Word 양식 파일의 구조를 분석한다.

    Args:
        file_data: Word 파일 바이너리 데이터

    Returns:
        양식 구조 딕셔너리 (template_structure 형식)

    Raises:
        ValueError: 파일을 읽을 수 없는 경우
    """
    try:
        doc = Document(io.BytesIO(file_data))
    except Exception as e:
        raise ValueError(f"Word 파일을 읽을 수 없습니다: {e}") from e

    # 1. 본문 단락에서 플레이스홀더 추출
    placeholders = _extract_placeholders_from_paragraphs(doc)

    # 2. 표 구조 분석
    tables = _analyze_tables(doc)

    # 3. 표 내부 플레이스홀더도 수집
    table_placeholders = _extract_placeholders_from_tables(doc)
    all_placeholders = list(dict.fromkeys(placeholders + table_placeholders))

    logger.info(
        "Word 양식 분석 완료: %d개 플레이스홀더, %d개 표",
        len(all_placeholders),
        len(tables),
    )

    return {
        "file_type": "docx",
        "sheets": [],
        "placeholders": all_placeholders,
        "tables": tables,
    }


def _extract_placeholders_from_paragraphs(doc: Document) -> list[str]:
    """문서 본문 단락에서 {{placeholder}} 패턴을 추출한다.

    Run이 분리되어 있는 경우를 처리하기 위해 단락 전체 텍스트에서 패턴을 찾는다.

    Args:
        doc: python-docx Document 객체

    Returns:
        플레이스홀더 이름 목록 ({{ }} 제거)
    """
    placeholders: list[str] = []

    for para in doc.paragraphs:
        text = para.text
        matches = _PLACEHOLDER_PATTERN.findall(text)
        for match in matches:
            name = match.strip()
            if name and name not in placeholders:
                placeholders.append(name)

    return placeholders


def _extract_placeholders_from_tables(doc: Document) -> list[str]:
    """표 내부에서 {{placeholder}} 패턴을 추출한다.

    Args:
        doc: python-docx Document 객체

    Returns:
        플레이스홀더 이름 목록
    """
    placeholders: list[str] = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                matches = _PLACEHOLDER_PATTERN.findall(text)
                for match in matches:
                    name = match.strip()
                    if name and name not in placeholders:
                        placeholders.append(name)

    return placeholders


def _analyze_tables(doc: Document) -> list[dict[str, Any]]:
    """문서 내 표 구조를 분석한다.

    각 표의 첫 번째 행을 헤더로 인식하고, 데이터 행 수와 플레이스홀더 존재 여부를 확인한다.

    Args:
        doc: python-docx Document 객체

    Returns:
        표 구조 목록
    """
    tables_info: list[dict[str, Any]] = []

    for idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        # 첫 번째 행을 헤더로 간주
        header_row = table.rows[0]
        headers: list[str] = []
        for cell in header_row.cells:
            text = cell.text.strip()
            headers.append(text)

        # 중복 헤더 제거 (병합 셀로 인한 중복)
        seen: set[str] = set()
        unique_headers: list[str] = []
        for h in headers:
            if h not in seen:
                unique_headers.append(h)
                seen.add(h)

        # 데이터 행 수 (헤더 행 제외)
        data_row_count = len(table.rows) - 1

        # 표 내부 플레이스홀더 존재 여부
        has_placeholder = False
        for row in table.rows[1:]:
            for cell in row.cells:
                if _PLACEHOLDER_PATTERN.search(cell.text):
                    has_placeholder = True
                    break
            if has_placeholder:
                break

        tables_info.append({
            "index": idx,
            "headers": unique_headers,
            "row_count": data_row_count,
            "has_placeholder_cells": has_placeholder,
        })

    return tables_info
