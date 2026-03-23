"""Word 양식 데이터 채우기 모듈.

python-docx를 사용하여 조회 결과를 Word 양식에 채워넣고,
원본의 스타일과 서식을 보존한다.
"""

from __future__ import annotations

import copy
import io
import logging
import re
from typing import Any, Optional
from xml.etree.ElementTree import Element

from docx import Document
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = logging.getLogger(__name__)

_PLACEHOLDER_PATTERN = re.compile(r"\{\{(.+?)\}\}")


def fill_word_template(
    file_data: bytes,
    template_structure: dict[str, Any],
    column_mapping: dict[str, Optional[str]],
    rows: list[dict[str, Any]],
    single_row: dict[str, Any] | None = None,
) -> bytes:
    """Word 양식에 조회 결과를 채워넣는다.

    Args:
        file_data: 원본 Word 파일 바이너리
        template_structure: 양식 구조 정보 (word_parser 출력)
        column_mapping: 필드-컬럼 매핑 (field_mapper 출력)
        rows: 조회 결과 행 목록 (표 채우기용)
        single_row: 단일 행 데이터 (플레이스홀더 치환용, 없으면 rows[0] 사용)

    Returns:
        데이터가 채워진 Word 파일 바이너리

    Raises:
        ValueError: 파일을 처리할 수 없는 경우
    """
    try:
        doc = Document(io.BytesIO(file_data))
    except Exception as e:
        raise ValueError(f"Word 파일을 읽을 수 없습니다: {e}") from e

    # 단일 행 데이터 결정
    fill_row = single_row or (rows[0] if rows else {})

    # 1. 본문 플레이스홀더 치환
    _replace_paragraph_placeholders(doc, column_mapping, fill_row)

    # 2. 표 데이터 채우기
    tables_info = template_structure.get("tables", [])
    _fill_tables(doc, tables_info, column_mapping, rows)

    # 바이너리로 저장
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    logger.info("Word 파일 생성 완료: %d건 데이터 채움", len(rows))
    return output.getvalue()


def _replace_paragraph_placeholders(
    doc: Document,
    column_mapping: dict[str, Optional[str]],
    data_row: dict[str, Any],
) -> None:
    """문서 본문 단락의 {{placeholder}}를 실제 데이터로 치환한다.

    Run 레벨에서 치환하여 스타일을 보존한다.

    Args:
        doc: python-docx Document 객체
        column_mapping: 필드-컬럼 매핑
        data_row: 단일 행 데이터
    """
    for para in doc.paragraphs:
        _replace_placeholders_in_paragraph(para, column_mapping, data_row)


def _replace_placeholders_in_paragraph(
    para: Paragraph,
    column_mapping: dict[str, Optional[str]],
    data_row: dict[str, Any],
) -> None:
    """단일 단락에서 플레이스홀더를 치환한다.

    {{placeholder}}가 여러 Run에 걸쳐 분리될 수 있으므로,
    먼저 전체 텍스트에서 패턴을 확인한 후 Run을 재조합한다.

    Args:
        para: python-docx Paragraph 객체
        column_mapping: 필드-컬럼 매핑
        data_row: 데이터 행
    """
    full_text = para.text
    if "{{" not in full_text:
        return

    matches = _PLACEHOLDER_PATTERN.findall(full_text)
    if not matches:
        return

    # Run 재조합: 전체 텍스트를 하나로 합쳐서 치환 후 첫 번째 run에 설정
    # 나머지 run은 비운다
    new_text = full_text
    for field_name in matches:
        field_name_stripped = field_name.strip()
        db_column = column_mapping.get(field_name_stripped)
        if db_column:
            value = _get_value_from_row(data_row, db_column)
            if value is None:
                # 매핑된 컬럼이지만 데이터에 값이 없음: 빈 문자열로 치환
                replacement = ""
            else:
                replacement = str(value)
        else:
            # 매핑되지 않은 필드 (column_mapping에서 None): 빈 문자열로 치환
            replacement = ""

        new_text = new_text.replace("{{" + field_name + "}}", replacement)

    # Run에 텍스트 반영 (스타일 보존)
    if para.runs:
        # 첫 번째 run에 전체 치환된 텍스트 설정
        para.runs[0].text = new_text
        # 나머지 run 비우기
        for run in para.runs[1:]:
            run.text = ""


def _fill_tables(
    doc: Document,
    tables_info: list[dict[str, Any]],
    column_mapping: dict[str, Optional[str]],
    rows: list[dict[str, Any]],
) -> None:
    """문서 내 표에 데이터를 채운다.

    Args:
        doc: python-docx Document 객체
        tables_info: 표 구조 정보 목록
        column_mapping: 필드-컬럼 매핑
        rows: 조회 결과 행 목록
    """
    for table_info in tables_info:
        table_idx = table_info.get("index", 0)
        if table_idx >= len(doc.tables):
            logger.warning("표 인덱스 %d가 범위를 초과하여 스킵", table_idx)
            continue

        table = doc.tables[table_idx]
        headers = table_info.get("headers", [])
        has_placeholders = table_info.get("has_placeholder_cells", False)

        if has_placeholders:
            # 플레이스홀더가 있는 표: 기존 셀의 플레이스홀더를 치환
            _fill_table_placeholders(table, column_mapping, rows)
        else:
            # 일반 표: 헤더 기반으로 데이터 행 추가
            _fill_table_rows(table, headers, column_mapping, rows)


def _fill_table_placeholders(
    table: Table,
    column_mapping: dict[str, Optional[str]],
    rows: list[dict[str, Any]],
) -> None:
    """표 내부의 플레이스홀더를 데이터로 치환한다.

    Args:
        table: python-docx Table 객체
        column_mapping: 필드-컬럼 매핑
        rows: 데이터 행 목록
    """
    data_row = rows[0] if rows else {}

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_placeholders_in_paragraph(para, column_mapping, data_row)


def _fill_table_rows(
    table: Table,
    headers: list[str],
    column_mapping: dict[str, Optional[str]],
    rows: list[dict[str, Any]],
) -> None:
    """표에 데이터 행을 추가한다.

    기존 빈 행에 먼저 채우고, 데이터가 더 있으면 새 행을 추가한다.

    Args:
        table: python-docx Table 객체
        headers: 표 헤더 목록
        column_mapping: 필드-컬럼 매핑
        rows: 데이터 행 목록
    """
    if not rows or not headers:
        return

    # 헤더별 매핑된 DB 컬럼 확인
    col_assignments: list[tuple[int, Optional[str]]] = []
    for idx, header in enumerate(headers):
        mapped = column_mapping.get(header)
        col_assignments.append((idx, mapped))

    # 기존 데이터 행 (헤더 행 = index 0 이후)
    existing_data_rows = list(table.rows[1:])

    # 스타일 참조용 행 (첫 번째 데이터 행 또는 헤더 행)
    style_ref_row = existing_data_rows[0] if existing_data_rows else table.rows[0]

    for row_idx, data_row in enumerate(rows):
        if row_idx < len(existing_data_rows):
            # 기존 행에 채우기
            target_row = existing_data_rows[row_idx]
        else:
            # 새 행 추가
            target_row = _add_row_with_style(table, style_ref_row)

        _fill_row_cells(target_row, col_assignments, data_row)


def _fill_row_cells(
    row: _Row,
    col_assignments: list[tuple[int, Optional[str]]],
    data_row: dict[str, Any],
) -> None:
    """행의 셀에 데이터를 채운다.

    Args:
        row: python-docx 행 객체
        col_assignments: (셀 인덱스, DB 컬럼) 목록
        data_row: 데이터 행
    """
    cells = list(row.cells)

    for cell_idx, db_column in col_assignments:
        if cell_idx >= len(cells):
            break
        if db_column is None:
            continue

        value = _get_value_from_row(data_row, db_column)
        cell = cells[cell_idx]

        # 셀 텍스트 설정 (기존 스타일 보존)
        if cell.paragraphs:
            para = cell.paragraphs[0]
            if para.runs:
                para.runs[0].text = str(value) if value is not None else ""
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = str(value) if value is not None else ""
        else:
            cell.text = str(value) if value is not None else ""


def _add_row_with_style(table: Table, ref_row: _Row) -> _Row:
    """참조 행의 스타일을 복사하여 새 행을 추가한다.

    Args:
        table: python-docx Table 객체
        ref_row: 스타일 참조 행

    Returns:
        추가된 새 행
    """
    # XML 레벨에서 행 복사
    new_tr = copy.deepcopy(ref_row._tr)

    # 셀 텍스트 비우기
    for tc in new_tr.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
        tc.text = ""

    table._tbl.append(new_tr)

    return table.rows[-1]


def _get_value_from_row(
    data_row: dict[str, Any],
    db_column: str,
) -> Any:
    """데이터 행에서 DB 컬럼에 해당하는 값을 추출한다.

    Args:
        data_row: 조회 결과 행
        db_column: "table.column" 형식의 DB 컬럼명

    Returns:
        추출된 값 또는 None
    """
    if db_column in data_row:
        return data_row[db_column]

    if "." in db_column:
        col_name = db_column.split(".", 1)[1]
        if col_name in data_row:
            return data_row[col_name]

    lower_col = db_column.lower()
    for key, value in data_row.items():
        if key.lower() == lower_col or (
            "." in db_column and key.lower() == db_column.split(".", 1)[1].lower()
        ):
            return value

    return None
