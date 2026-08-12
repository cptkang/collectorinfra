"""DRM 연동 진단 (Plan 69 §4.2 — 어드민 진단 도구).

실기 환경이 운영계뿐이므로 셸 없이 브라우저에서 연동 상태를 점검할 수 있게 한다.

**복호화 결과 파일은 반환하지 않는다** — 진단 정보(ret 값, 시그니처, 파싱 가능
여부)만 돌려주므로 반복 호출로도 문서 내용이 복원되지 않는다(복호화 오라클 방지).

두 함수 모두 예외를 던지지 않고 **항상 구조화된 결과를 반환**한다. 실패가 곧
진단 데이터이며, 화면이 에러로 깨지면 정작 필요한 ret 값을 볼 수 없기 때문이다.
"""

from __future__ import annotations

import asyncio
import io
import logging
import os
import subprocess
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING, Any, Optional

from src.infrastructure.drm.decryptor import DrmDecryptError, get_decryptor
from src.infrastructure.drm.detector import detect_file_kind, header_hex

if TYPE_CHECKING:
    from src.config import DrmConfig

logger = logging.getLogger(__name__)

# KeyManager는 24시간 주기로 키 파일을 갱신한다(가이드: 키갱신 매 24시간 필수).
# 여유 6시간을 둬 30시간을 초과하면 KeyManager 이상으로 판정한다.
_KEY_STALE_HOURS = 30

_JAVA_VERSION_TIMEOUT_SEC = 5


def _check_path(label: str, raw: str) -> dict[str, Any]:
    """설정 경로 1건의 존재·읽기 가능 여부를 판정한다."""
    result: dict[str, Any] = {
        "label": label,
        "path": raw,
        "configured": bool(raw),
        "exists": False,
        "readable": False,
    }
    if not raw:
        result["message"] = "설정되지 않음"
        return result

    path = Path(raw)
    result["exists"] = path.exists()
    if not result["exists"]:
        result["message"] = "파일이 존재하지 않음"
        return result

    result["readable"] = os.access(path, os.R_OK)
    if not result["readable"]:
        result["message"] = "읽기 권한 없음"
        return result

    try:
        result["size_bytes"] = path.stat().st_size
    except OSError as e:  # pragma: no cover - 경합 상황
        result["message"] = f"stat 실패: {e}"
        return result

    result["message"] = "정상"
    return result


def _check_key_file(raw: str) -> dict[str, Any]:
    """키 파일 상태를 점검한다.

    키 파일 mtime은 KeyManager 생존 신호다 — 24시간 주기 갱신이 멈추면
    ret 3000/3003/3030 계열 실패로 이어지므로, 경과 시간을 함께 보고한다.
    """
    result = _check_path("키 파일 (keyDAC_SVR0.sc)", raw)
    if not result.get("exists"):
        return result

    try:
        mtime = datetime.fromtimestamp(Path(raw).stat().st_mtime, tz=timezone.utc)
    except OSError as e:  # pragma: no cover - 경합 상황
        result["message"] = f"mtime 조회 실패: {e}"
        return result

    age_hours = (datetime.now(timezone.utc) - mtime).total_seconds() / 3600
    result["modified_at"] = mtime.isoformat()
    result["age_hours"] = round(age_hours, 1)
    result["stale"] = age_hours > _KEY_STALE_HOURS

    if result["stale"]:
        result["message"] = (
            f"키 갱신이 {result['age_hours']}시간째 없음 — "
            "KeyManager 동작 확인 필요(갱신 주기 24시간)"
        )
    elif result.get("readable"):
        result["message"] = f"정상 (최근 갱신 {result['age_hours']}시간 전)"
    return result


def _check_java(java_bin: str) -> dict[str, Any]:
    """java 실행 파일 가용성과 버전을 확인한다 (요구: 1.8 이상)."""
    result: dict[str, Any] = {"label": "Java 런타임", "command": java_bin}
    try:
        proc = subprocess.run(
            [java_bin, "-version"],
            capture_output=True,
            timeout=_JAVA_VERSION_TIMEOUT_SEC,
        )
    except FileNotFoundError:
        result["available"] = False
        result["message"] = f"java 실행 파일을 찾을 수 없음: {java_bin}"
        return result
    except subprocess.TimeoutExpired:
        result["available"] = False
        result["message"] = "java -version 응답 없음(타임아웃)"
        return result
    except OSError as e:
        result["available"] = False
        result["message"] = f"java 실행 실패: {e}"
        return result

    # java -version은 관례적으로 stderr에 출력한다
    raw = (proc.stderr or proc.stdout or b"").decode("utf-8", errors="replace")
    first_line = raw.strip().splitlines()[0] if raw.strip() else ""
    result["available"] = proc.returncode == 0
    result["version"] = first_line
    result["message"] = first_line or "버전 정보를 확인할 수 없음"
    return result


def _check_temp_dir(raw: str) -> dict[str, Any]:
    """복호화 작업 디렉터리 상태를 확인한다(빈 값이면 시스템 temp 하위 자동 생성)."""
    import tempfile

    path = Path(raw) if raw else Path(tempfile.gettempdir()) / "drm_scsl"
    result: dict[str, Any] = {
        "label": "작업 디렉터리",
        "path": str(path),
        "configured": bool(raw),
        "exists": path.exists(),
    }
    if not result["exists"]:
        result["message"] = "미생성 (첫 복호화 시 자동 생성)"
        return result

    result["writable"] = os.access(path, os.W_OK)
    try:
        result["leftover_files"] = len(list(path.glob("*")))
    except OSError:
        result["leftover_files"] = None
    result["message"] = "정상" if result["writable"] else "쓰기 권한 없음"
    return result


def check_environment(config: "DrmConfig") -> dict[str, Any]:
    """DRM 연동 환경 상태를 점검한다 (파일 업로드 없음).

    Returns:
        enabled 여부, 경로 4종·키 파일·java·temp 점검 결과, 종합 판정
    """
    checks = [
        _check_path("래퍼 (Decrypt.java)", config.wrapper_path),
        _check_path("scsl.jar", config.scsl_jar_path),
        _check_path("softcamp.properties", config.properties_path),
        _check_key_file(config.key_file_path),
    ]
    java = _check_java(config.java_bin)
    temp = _check_temp_dir(config.temp_dir)

    if not config.enabled:
        summary = "비활성 (DRM_ENABLED=false — 업로드 파일을 그대로 처리)"
        ready = False
    else:
        problems = [c for c in checks if not (c.get("exists") and c.get("readable"))]
        stale_key = any(c.get("stale") for c in checks)
        if problems:
            summary = f"설정 문제 {len(problems)}건 — 아래 항목 확인 필요"
            ready = False
        elif not java.get("available"):
            summary = "Java 런타임 확인 필요"
            ready = False
        elif stale_key:
            summary = "키 파일 갱신 지연 — KeyManager 동작 확인 필요"
            ready = False
        else:
            summary = "정상 — 복호화 준비 완료"
            ready = True

    return {
        "enabled": config.enabled,
        "ready": ready,
        "summary": summary,
        "group_id": config.group_id,
        "timeout_sec": config.timeout_sec,
        "checks": checks,
        "java": java,
        "temp_dir": temp,
    }


def _inspect_plain(plain: bytes, filename: str) -> dict[str, Any]:
    """복호화 산출물이 폼필이 처리 가능한 파일인지 확인한다.

    시트명·문단 수 수준까지만 보고하고 **셀 값은 반환하지 않는다**.
    """
    result: dict[str, Any] = {
        "size_bytes": len(plain),
        "is_zip": plain.startswith(b"PK\x03\x04"),
        "header_hex": header_hex(plain),
    }
    ext = Path(filename).suffix.lower()

    try:
        if ext == ".xlsx":
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(plain), read_only=True)
            result["parsed"] = True
            result["sheet_names"] = list(wb.sheetnames)[:20]
            wb.close()
        elif ext == ".docx":
            import docx

            document = docx.Document(io.BytesIO(plain))
            result["parsed"] = True
            result["paragraph_count"] = len(document.paragraphs)
            result["table_count"] = len(document.tables)
        else:
            result["parsed"] = None
            result["parse_message"] = f"파싱 검증 미지원 확장자: {ext or '(없음)'}"
            return result
        result["parse_message"] = "폼필이 처리 가능한 파일로 확인됨"
    except Exception as e:
        result["parsed"] = False
        result["parse_message"] = f"복호화는 됐으나 문서 파싱 실패: {e}"

    return result


async def verify_sample(
    config: "DrmConfig", file_bytes: bytes, filename: str
) -> dict[str, Any]:
    """암호화 샘플을 복호화해 진단 결과만 반환한다 (평문 미반환).

    예외를 던지지 않고 항상 결과 dict를 반환한다 — 실패도 진단 데이터다.
    """
    kind = detect_file_kind(file_bytes)
    result: dict[str, Any] = {
        "file_name": filename,
        "file_size_bytes": len(file_bytes),
        "detected": kind,
        "header_hex": header_hex(file_bytes),
        "success": False,
        "ret": None,
        "elapsed_ms": None,
    }

    if kind == "unknown":
        result["message"] = (
            "DRM 암호문도 평문 문서(ZIP)도 아닙니다 — 손상 파일이거나 지원하지 "
            "않는 형식입니다."
        )
        return result

    if kind == "plain":
        result["message"] = (
            "평문 문서(ZIP)입니다 — 복호화 대상이 아닙니다. DRM이 적용된 샘플을 "
            "업로드하세요(문서를 한 번 열었다 닫으면 암호화됩니다)."
        )
        result["output"] = _inspect_plain(file_bytes, filename)
        return result

    if not config.enabled:
        result["message"] = (
            "DRM 암호문으로 확인됐으나 DRM_ENABLED=false 상태라 복호화를 수행하지 "
            "않았습니다. 감지는 정상 동작합니다."
        )
        return result

    decryptor = get_decryptor(config)
    start = time.time()
    try:
        # 진단은 ret 원시값이 필요하다(0=복호화 성공 / -36=원본이 애초에 평문).
        detailed = getattr(decryptor, "decrypt_detailed", None)
        if detailed is not None:
            plain, ret = await detailed(file_bytes, filename)
        else:
            plain, ret = await decryptor.decrypt(file_bytes, filename), 0
    except DrmDecryptError as e:
        result["elapsed_ms"] = round((time.time() - start) * 1000, 1)
        result["ret"] = e.ret_code
        result["message"] = f"복호화 실패: {e.reason}"
        result["detail"] = e.detail
        logger.warning(
            "DRM 진단 복호화 실패: file=%s reason=%s ret=%s", filename, e.reason, e.ret_code
        )
        return result
    except Exception as e:  # 진단 도구는 어떤 예외에서도 화면이 깨지지 않아야 한다
        result["elapsed_ms"] = round((time.time() - start) * 1000, 1)
        result["message"] = f"복호화 중 예기치 못한 오류: {e}"
        logger.exception("DRM 진단 중 예외: file=%s", filename)
        return result

    result["elapsed_ms"] = round((time.time() - start) * 1000, 1)
    result["ret"] = ret
    result["success"] = True
    result["output"] = _inspect_plain(plain, filename)
    # 평문 바이트는 반환하지 않는다 — 참조를 즉시 끊어 GC 대상으로 만든다
    del plain

    # -36은 파싱 실패의 근본 원인이 될 수 있으므로 먼저 보고한다:
    # 감지는 SCDS 암호문이었는데 scsl이 평문으로 판정했다면 복호화가 실제로
    # 수행되지 않은 것이다(키·정책 설정 이상).
    if ret == -36:
        result["message"] = (
            "scsl이 원본을 평문으로 판정했습니다(ret -36) — 감지는 암호문이었으므로 "
            "키·정책 설정을 확인하세요."
        )
    elif result["output"].get("parsed") is False:
        result["message"] = "복호화는 성공했으나 문서 파싱에 실패했습니다"
    else:
        result["message"] = "복호화 성공"
    return result
