"""문서 처리 통합 테스트.

파싱 -> 매핑 -> 생성 전체 흐름을 검증한다.
"""

from __future__ import annotations

import io
import json
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from src.document.excel_parser import parse_excel_template
from src.document.excel_writer import fill_excel_template
from src.document.field_mapper import map_fields
from src.document.word_parser import parse_word_template
from src.document.word_writer import fill_word_template


def _create_excel_template() -> bytes:
    """통합 테스트용 Excel 양식을 생성한다."""
    wb = Workbook()
    ws = wb.active
    ws.title = "서버 현황"

    headers = ["서버명", "IP주소", "CPU 사용률(%)", "메모리(GB)", "비고"]
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = cell.font.copy(bold=True)

    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    return buf.getvalue()


def _create_word_template() -> bytes:
    """통합 테스트용 Word 양식을 생성한다."""
    doc = Document()
    doc.add_paragraph("서버 점검 보고서")
    doc.add_paragraph("작성일: {{작성일}}")
    doc.add_paragraph("작성자: {{작성자}}")
    doc.add_paragraph("")

    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "서버명"
    table.rows[0].cells[1].text = "IP"
    table.rows[0].cells[2].text = "상태"
    table.rows[1].cells[0].text = ""
    table.rows[1].cells[1].text = ""
    table.rows[1].cells[2].text = ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_schema() -> dict:
    """테스트용 스키마."""
    return {
        "tables": {
            "servers": {
                "columns": [
                    {"name": "hostname", "type": "varchar(255)"},
                    {"name": "ip_address", "type": "varchar(50)"},
                    {"name": "status", "type": "varchar(20)"},
                ],
            },
            "cpu_metrics": {
                "columns": [
                    {"name": "server_id", "type": "integer"},
                    {"name": "usage_pct", "type": "float"},
                ],
            },
            "memory_metrics": {
                "columns": [
                    {"name": "server_id", "type": "integer"},
                    {"name": "total_gb", "type": "float"},
                ],
            },
        }
    }


class TestExcelIntegration:
    """Excel 파싱 -> 매핑 -> 생성 통합 테스트."""

    @pytest.mark.asyncio
    async def test_full_excel_flow(self):
        """Excel 양식 전체 처리 흐름."""
        # 1. 양식 생성
        template_bytes = _create_excel_template()

        # 2. 파싱
        template_structure = parse_excel_template(template_bytes)
        assert len(template_structure["sheets"]) == 1
        assert "서버명" in template_structure["sheets"][0]["headers"]

        # 3. LLM 매핑 (mock)
        mock_llm = AsyncMock()
        mock_llm.ainvoke.return_value = MagicMock(
            content=json.dumps({
                "서버명": "servers.hostname",
                "IP주소": "servers.ip_address",
                "CPU 사용률(%)": "cpu_metrics.usage_pct",
                "메모리(GB)": "memory_metrics.total_gb",
                "비고": None,
            })
        )

        schema = _make_schema()
        mapping = await map_fields(mock_llm, template_structure, schema)

        assert mapping["서버명"] == "servers.hostname"
        assert mapping["비고"] is None

        # 4. 데이터 채우기
        rows = [
            {"hostname": "web-01", "ip_address": "10.0.0.1", "usage_pct": 85.2, "total_gb": 32.0},
            {"hostname": "web-02", "ip_address": "10.0.0.2", "usage_pct": 62.1, "total_gb": 16.0},
            {"hostname": "db-01", "ip_address": "10.0.1.1", "usage_pct": 92.5, "total_gb": 64.0},
        ]

        result_bytes, _ = fill_excel_template(
            template_bytes, template_structure, mapping, rows
        )

        # 5. 결과 검증
        wb = load_workbook(io.BytesIO(result_bytes))
        ws = wb.active

        # 헤더 보존
        assert ws.cell(row=1, column=1).value == "서버명"

        # 데이터 검증
        assert ws.cell(row=2, column=1).value == "web-01"
        assert ws.cell(row=2, column=2).value == "10.0.0.1"
        assert ws.cell(row=2, column=3).value == 85.2
        assert ws.cell(row=2, column=4).value == 32.0
        assert ws.cell(row=2, column=5).value is None  # 비고는 매핑 없음

        assert ws.cell(row=4, column=1).value == "db-01"
        assert ws.cell(row=4, column=3).value == 92.5

        wb.close()


class TestWordIntegration:
    """Word 파싱 -> 매핑 -> 생성 통합 테스트."""

    @pytest.mark.asyncio
    async def test_full_word_flow(self):
        """Word 양식 전체 처리 흐름."""
        # 1. 양식 생성
        template_bytes = _create_word_template()

        # 2. 파싱
        template_structure = parse_word_template(template_bytes)
        assert "작성일" in template_structure["placeholders"]
        assert "작성자" in template_structure["placeholders"]
        assert len(template_structure["tables"]) == 1

        # 3. LLM 매핑 (mock)
        mock_llm = AsyncMock()
        mock_llm.ainvoke.return_value = MagicMock(
            content=json.dumps({
                "작성일": None,
                "작성자": None,
                "서버명": "servers.hostname",
                "IP": "servers.ip_address",
                "상태": "servers.status",
            })
        )

        schema = _make_schema()
        mapping = await map_fields(mock_llm, template_structure, schema)

        assert mapping["서버명"] == "servers.hostname"
        assert mapping["작성일"] is None  # DB 매핑 없는 필드

        # 4. 데이터 채우기
        rows = [
            {"hostname": "web-01", "ip_address": "10.0.0.1", "status": "정상"},
            {"hostname": "web-02", "ip_address": "10.0.0.2", "status": "점검중"},
        ]

        # 매핑에 없는 필드는 직접 값 지정
        mapping_with_manual = {**mapping, "작성일": None, "작성자": None}

        result_bytes = fill_word_template(
            template_bytes,
            template_structure,
            mapping_with_manual,
            rows,
            single_row={"hostname": "web-01"},  # placeholder용
        )

        # 5. 결과 검증
        doc = Document(io.BytesIO(result_bytes))

        # 표 데이터 검증
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "서버명"  # 헤더 보존
        assert table.rows[1].cells[0].text == "web-01"
        assert table.rows[1].cells[1].text == "10.0.0.1"
        assert table.rows[1].cells[2].text == "정상"

        # 두 번째 행
        assert len(table.rows) >= 3
        assert table.rows[2].cells[0].text == "web-02"
