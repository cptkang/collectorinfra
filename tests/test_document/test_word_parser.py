"""Word 양식 파서 단위 테스트."""

from __future__ import annotations

import io

import pytest
from docx import Document

from src.document.word_parser import parse_word_template


def _create_word_bytes(
    paragraphs: list[str] | None = None,
    tables: list[dict] | None = None,
) -> bytes:
    """테스트용 Word 파일 바이너리를 생성한다.

    Args:
        paragraphs: 단락 텍스트 목록
        tables: 표 정보 목록 [{"headers": [...], "rows": [[...]]}]
    """
    doc = Document()

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    if tables:
        for table_info in tables:
            headers = table_info.get("headers", [])
            rows_data = table_info.get("rows", [])
            total_rows = 1 + len(rows_data)
            cols = len(headers) if headers else 1

            table = doc.add_table(rows=total_rows, cols=cols)
            # 헤더 행
            for col_idx, header in enumerate(headers):
                table.rows[0].cells[col_idx].text = header
            # 데이터 행
            for row_idx, row_data in enumerate(rows_data, 1):
                for col_idx, value in enumerate(row_data):
                    table.rows[row_idx].cells[col_idx].text = str(value)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestParseWordTemplate:
    """parse_word_template 함수 테스트."""

    def test_basic_placeholders(self):
        """본문 플레이스홀더를 추출한다."""
        data = _create_word_bytes(
            paragraphs=[
                "서버 점검 보고서",
                "서버명: {{서버명}}",
                "IP 주소: {{IP주소}}",
                "점검일: {{날짜}}",
            ]
        )

        result = parse_word_template(data)

        assert result["file_type"] == "docx"
        assert "서버명" in result["placeholders"]
        assert "IP주소" in result["placeholders"]
        assert "날짜" in result["placeholders"]
        assert result["sheets"] == []

    def test_table_structure(self):
        """표 구조를 분석한다."""
        data = _create_word_bytes(
            tables=[{
                "headers": ["서버명", "IP", "상태"],
                "rows": [
                    ["web-01", "10.0.0.1", "정상"],
                    ["web-02", "10.0.0.2", "점검중"],
                ],
            }]
        )

        result = parse_word_template(data)

        assert len(result["tables"]) == 1
        table = result["tables"][0]
        assert table["index"] == 0
        assert table["headers"] == ["서버명", "IP", "상태"]
        assert table["row_count"] == 2

    def test_table_with_placeholders(self):
        """표 내부 플레이스홀더를 탐지한다."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "서버명"
        table.rows[0].cells[1].text = "상태"
        table.rows[1].cells[0].text = "{{서버명}}"
        table.rows[1].cells[1].text = "{{상태}}"

        buf = io.BytesIO()
        doc.save(buf)

        result = parse_word_template(buf.getvalue())

        assert len(result["tables"]) == 1
        assert result["tables"][0]["has_placeholder_cells"] is True
        assert "서버명" in result["placeholders"]
        assert "상태" in result["placeholders"]

    def test_no_placeholders_or_tables(self):
        """플레이스홀더나 표가 없는 문서."""
        data = _create_word_bytes(
            paragraphs=["일반 텍스트 문서입니다."]
        )

        result = parse_word_template(data)

        assert result["placeholders"] == []
        assert result["tables"] == []

    def test_duplicate_placeholders_deduplicated(self):
        """중복 플레이스홀더는 제거된다."""
        data = _create_word_bytes(
            paragraphs=[
                "{{서버명}}은 {{서버명}}입니다.",
                "IP: {{IP}}",
            ]
        )

        result = parse_word_template(data)

        # 서버명이 2번 나와도 1번만 나와야 함
        assert result["placeholders"].count("서버명") == 1
        assert "IP" in result["placeholders"]

    def test_multiple_tables(self):
        """다중 표 파싱."""
        data = _create_word_bytes(
            tables=[
                {"headers": ["이름", "값"], "rows": [["a", "1"]]},
                {"headers": ["항목", "수량", "단가"], "rows": []},
            ]
        )

        result = parse_word_template(data)

        assert len(result["tables"]) == 2
        assert result["tables"][0]["index"] == 0
        assert result["tables"][1]["index"] == 1
        assert len(result["tables"][1]["headers"]) == 3

    def test_invalid_file_raises_error(self):
        """유효하지 않은 파일은 ValueError를 발생시킨다."""
        with pytest.raises(ValueError, match="Word 파일을 읽을 수 없습니다"):
            parse_word_template(b"not a valid docx file")

    def test_mixed_paragraphs_and_tables(self):
        """본문 플레이스홀더와 표가 함께 있는 문서."""
        data = _create_word_bytes(
            paragraphs=[
                "보고서 제목: {{제목}}",
                "작성일: {{날짜}}",
            ],
            tables=[{
                "headers": ["서버명", "CPU", "메모리"],
                "rows": [],
            }],
        )

        result = parse_word_template(data)

        assert "제목" in result["placeholders"]
        assert "날짜" in result["placeholders"]
        assert len(result["tables"]) == 1
        assert result["tables"][0]["headers"] == ["서버명", "CPU", "메모리"]
