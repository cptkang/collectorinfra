"""Word 양식 데이터 채우기 단위 테스트."""

from __future__ import annotations

import io

import pytest
from docx import Document

from src.document.word_parser import parse_word_template
from src.document.word_writer import fill_word_template


def _create_word_template_bytes(
    paragraphs: list[str] | None = None,
    tables: list[dict] | None = None,
) -> bytes:
    """테스트용 Word 양식 파일을 생성한다."""
    doc = Document()

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    if tables:
        for table_info in tables:
            headers = table_info.get("headers", [])
            rows_data = table_info.get("rows", [])
            total_rows = 1 + len(rows_data)
            cols = len(headers)

            table = doc.add_table(rows=total_rows, cols=cols)
            for col_idx, header in enumerate(headers):
                table.rows[0].cells[col_idx].text = header
            for row_idx, row_data in enumerate(rows_data, 1):
                for col_idx, value in enumerate(row_data):
                    table.rows[row_idx].cells[col_idx].text = str(value)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestFillWordTemplate:
    """fill_word_template 함수 테스트."""

    def test_placeholder_replacement(self):
        """본문 플레이스홀더가 데이터로 치환된다."""
        template_bytes = _create_word_template_bytes(
            paragraphs=[
                "서버 보고서",
                "서버명: {{서버명}}",
                "IP 주소: {{IP주소}}",
            ]
        )
        template_structure = parse_word_template(template_bytes)
        column_mapping = {
            "서버명": "servers.hostname",
            "IP주소": "servers.ip_address",
        }
        rows = [{"hostname": "web-01", "ip_address": "10.0.0.1"}]

        result_bytes = fill_word_template(
            template_bytes, template_structure, column_mapping, rows
        )

        doc = Document(io.BytesIO(result_bytes))
        texts = [p.text for p in doc.paragraphs]
        full_text = "\n".join(texts)

        assert "web-01" in full_text
        assert "10.0.0.1" in full_text
        assert "{{서버명}}" not in full_text
        assert "{{IP주소}}" not in full_text

    def test_table_data_fill(self):
        """표에 데이터 행이 채워진다."""
        template_bytes = _create_word_template_bytes(
            tables=[{
                "headers": ["서버명", "IP", "상태"],
                "rows": [["", "", ""]],  # 빈 데이터 행
            }]
        )
        template_structure = parse_word_template(template_bytes)
        column_mapping = {
            "서버명": "servers.hostname",
            "IP": "servers.ip_address",
            "상태": "servers.status",
        }
        rows = [
            {"hostname": "web-01", "ip_address": "10.0.0.1", "status": "active"},
            {"hostname": "web-02", "ip_address": "10.0.0.2", "status": "inactive"},
        ]

        result_bytes = fill_word_template(
            template_bytes, template_structure, column_mapping, rows
        )

        doc = Document(io.BytesIO(result_bytes))
        table = doc.tables[0]
        # 헤더 행 보존
        assert table.rows[0].cells[0].text == "서버명"
        # 첫 번째 데이터 행
        assert table.rows[1].cells[0].text == "web-01"
        assert table.rows[1].cells[1].text == "10.0.0.1"
        # 두 번째 데이터 행 (새로 추가됨)
        assert len(table.rows) >= 3
        assert table.rows[2].cells[0].text == "web-02"

    def test_table_placeholder_replacement(self):
        """표 내부의 플레이스홀더가 치환된다."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "서버명"
        table.rows[0].cells[1].text = "상태"
        table.rows[1].cells[0].text = "{{서버명}}"
        table.rows[1].cells[1].text = "{{상태}}"

        buf = io.BytesIO()
        doc.save(buf)
        template_bytes = buf.getvalue()

        template_structure = parse_word_template(template_bytes)
        column_mapping = {
            "서버명": "servers.hostname",
            "상태": "servers.status",
        }
        rows = [{"hostname": "web-01", "status": "active"}]

        result_bytes = fill_word_template(
            template_bytes, template_structure, column_mapping, rows
        )

        doc = Document(io.BytesIO(result_bytes))
        table = doc.tables[0]
        assert table.rows[1].cells[0].text == "web-01"
        assert table.rows[1].cells[1].text == "active"

    def test_unmapped_placeholder_cleared(self):
        """매핑되지 않은 플레이스홀더는 빈 문자열로 치환된다."""
        template_bytes = _create_word_template_bytes(
            paragraphs=["비고: {{비고}}"]
        )
        template_structure = parse_word_template(template_bytes)
        column_mapping = {"비고": None}
        rows = [{}]

        result_bytes = fill_word_template(
            template_bytes, template_structure, column_mapping, rows
        )

        doc = Document(io.BytesIO(result_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "{{비고}}" not in full_text

    def test_empty_rows(self):
        """빈 결과 행 처리."""
        template_bytes = _create_word_template_bytes(
            paragraphs=["서버명: {{서버명}}"]
        )
        template_structure = parse_word_template(template_bytes)
        column_mapping = {"서버명": "servers.hostname"}

        result_bytes = fill_word_template(
            template_bytes, template_structure, column_mapping, []
        )

        # 에러 없이 파일이 생성되어야 함
        doc = Document(io.BytesIO(result_bytes))
        assert len(doc.paragraphs) > 0

    def test_single_row_parameter(self):
        """single_row 매개변수로 플레이스홀더 데이터를 지정한다."""
        template_bytes = _create_word_template_bytes(
            paragraphs=["서버명: {{서버명}}"]
        )
        template_structure = parse_word_template(template_bytes)
        column_mapping = {"서버명": "servers.hostname"}

        result_bytes = fill_word_template(
            template_bytes,
            template_structure,
            column_mapping,
            rows=[],
            single_row={"hostname": "specific-server"},
        )

        doc = Document(io.BytesIO(result_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "specific-server" in full_text

    def test_invalid_file_raises_error(self):
        """유효하지 않은 파일은 ValueError를 발생시킨다."""
        with pytest.raises(ValueError):
            fill_word_template(
                b"invalid",
                {"tables": []},
                {},
                [],
            )

    def test_mixed_placeholders_and_tables(self):
        """플레이스홀더와 표가 함께 있는 문서."""
        template_bytes = _create_word_template_bytes(
            paragraphs=[
                "보고서: {{제목}}",
                "날짜: {{날짜}}",
            ],
            tables=[{
                "headers": ["서버명", "CPU"],
                "rows": [["", ""]],
            }],
        )
        template_structure = parse_word_template(template_bytes)
        column_mapping = {
            "제목": "reports.title",
            "날짜": "reports.date",
            "서버명": "servers.hostname",
            "CPU": "servers.cpu_usage",
        }
        rows = [{"hostname": "web-01", "cpu_usage": "85%"}]

        result_bytes = fill_word_template(
            template_bytes,
            template_structure,
            column_mapping,
            rows,
            single_row={"title": "월간 보고", "date": "2026-03-17"},
        )

        doc = Document(io.BytesIO(result_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "월간 보고" in full_text
        assert "2026-03-17" in full_text

        table = doc.tables[0]
        assert table.rows[1].cells[0].text == "web-01"
