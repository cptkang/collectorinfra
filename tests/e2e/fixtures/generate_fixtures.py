"""E2E 테스트용 fixture 파일(Excel, Word)을 생성하는 스크립트.

실행 방법:
    python tests/e2e/fixtures/generate_fixtures.py
"""

from __future__ import annotations

from pathlib import Path


def create_sample_xlsx(output_path: Path) -> None:
    """테스트용 Excel 양식을 생성한다.

    헤더: A1="서버명", B1="IP주소", C1="CPU사용률"
    2~6행은 비어있는 상태로 둔다.

    Args:
        output_path: 출력 파일 경로
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "서버현황"

    # 헤더 스타일
    header_font = Font(name="맑은 고딕", bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    headers = ["서버명", "IP주소", "CPU사용률"]
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # 2~6행 빈 셀에 테두리만 적용
    for row_idx in range(2, 7):
        for col_idx in range(1, 4):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = thin_border

    # 열 너비 설정
    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 15

    wb.save(output_path)
    print(f"  [OK] {output_path}")


def create_sample_docx(output_path: Path) -> None:
    """테스트용 Word 양식을 생성한다.

    본문: "서버 현황 보고서", {{서버명}}, {{IP주소}} 플레이스홀더
    테이블: 3열 헤더(서버명, IP주소, CPU사용률)

    Args:
        output_path: 출력 파일 경로
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()

    # 제목
    title = doc.add_heading("서버 현황 보고서", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 플레이스홀더가 포함된 본문
    doc.add_paragraph("")
    p1 = doc.add_paragraph()
    p1.add_run("서버명: ").bold = True
    p1.add_run("{{서버명}}")

    p2 = doc.add_paragraph()
    p2.add_run("IP주소: ").bold = True
    p2.add_run("{{IP주소}}")

    doc.add_paragraph("")

    # 테이블
    doc.add_heading("서버 목록", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # 헤더 행
    header_cells = table.rows[0].cells
    header_cells[0].text = "서버명"
    header_cells[1].text = "IP주소"
    header_cells[2].text = "CPU사용률"

    # 헤더 셀 굵게
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 빈 데이터 행 5개 추가
    for _ in range(5):
        table.add_row()

    doc.save(output_path)
    print(f"  [OK] {output_path}")


def main() -> None:
    """fixture 파일을 생성한다."""
    fixtures_dir = Path(__file__).parent

    print("E2E 테스트 fixture 파일 생성:")
    create_sample_xlsx(fixtures_dir / "sample_template.xlsx")
    create_sample_docx(fixtures_dir / "sample_template.docx")
    print("완료.")


if __name__ == "__main__":
    main()
